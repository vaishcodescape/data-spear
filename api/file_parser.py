"""
File format parsing for OmniGraph document ingestion.

Supported formats:
  - Plain text  (.txt)
  - PDF         (.pdf)   — requires pdfminer.six
  - DOCX        (.docx)  — requires python-docx
  - URL                  — requires httpx + beautifulsoup4
"""
import io
import logging
from typing import Optional

logger = logging.getLogger("omnigraph.file_parser")


def parse_pdf(content: bytes) -> str:
    """Extract text from a PDF byte stream using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text  # type: ignore
        text = extract_text(io.BytesIO(content))
        return text or ""
    except Exception as exc:
        logger.error("PDF parsing failed: %s", exc)
        raise ValueError(f"Could not parse PDF: {exc}") from exc


def parse_docx(content: bytes) -> str:
    """Extract text from a DOCX byte stream using python-docx."""
    try:
        from docx import Document  # type: ignore
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX parsing failed: %s", exc)
        raise ValueError(f"Could not parse DOCX: {exc}") from exc


def parse_url(url: str, timeout: int = 30) -> tuple[str, Optional[str]]:
    """
    Fetch a URL and return (text_content, page_title).
    Strips scripts, styles, nav, and footer elements.
    """
    try:
        import httpx
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise ImportError(
            "httpx and beautifulsoup4 are required for URL ingestion. "
            "Run: pip install httpx beautifulsoup4"
        ) from exc

    try:
        response = httpx.get(
            url,
            timeout=timeout,
            follow_redirects=True,
            headers={"User-Agent": "OmniGraph/1.0 (knowledge-graph-bot)"},
        )
        response.raise_for_status()
    except httpx.HTTPError as exc:
        raise ValueError(f"Failed to fetch URL '{url}': {exc}") from exc

    soup = BeautifulSoup(response.text, "html.parser")

    # Extract title before stripping tags
    title_tag = soup.find("title")
    page_title = title_tag.get_text(strip=True) if title_tag else None

    # Remove noise elements
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    # Collapse excessive blank lines
    import re
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip(), page_title


def parse_file(filename: str, content: bytes) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Falls back to UTF-8 decode for unknown types.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return parse_pdf(content)
    elif ext in ("docx",):
        return parse_docx(content)
    elif ext == "txt":
        return content.decode("utf-8", errors="replace")
    else:
        # Best-effort: try UTF-8
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="replace")
